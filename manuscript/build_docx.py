"""
manuscript/build_docx.py
========================
Render the Nano-FFI manuscript as Word documents, one per target journal.

The prose matches the humanized reading copy (manuscript.md); every reported
number is injected at build time from results.compute_results(), so the built
manuscripts cannot drift from the live extension. Figures are embedded from
manuscript/figures/ (run generate_figures.py first).

    python manuscript/generate_figures.py
    python manuscript/build_docx.py            # builds all journals
    python manuscript/build_docx.py Ingeniare  # or a single journal

Each journal differs only in page size, fonts, and margins; the research and the
numbers are identical, so only one venue may hold the paper at a time.
"""

import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from results import compute_results

FIG = os.path.join(HERE, "figures")

TITLE = ("Nano-FFI: a comptime-generated, branch-free foreign function "
         "interface between Python and Zig")
AUTHOR = "Felipe Carvajal Brown"
AFFIL = ("M.Sc. in Numerical Simulation in Engineering "
         "(Universidad Politecnica de Madrid, Spain); "
         "Independent Researcher, Santiago, Chile")
EMAIL = "fcarvajalbrown@gmail.com"

# Each profile: page size, margins, and body font. US Letter for the
# US/Elsevier venues, A4 for Ingeniare (Chile).
JOURNALS = {
    "Ingeniare": dict(folder="Ingeniare", font="Times New Roman", size=10,
                      page=("A4",), margin_cm=2.5),
    "PeerJ_CS": dict(folder="PeerJ_CS", font="Times New Roman", size=11,
                     page=("Letter",), margin_cm=2.54),
    "SoftwareX": dict(folder="SoftwareX", font="Times New Roman", size=10,
                      page=("Letter",), margin_cm=2.54),
    "Software_Impacts": dict(folder="Software_Impacts", font="Times New Roman",
                             size=10, page=("Letter",), margin_cm=2.54),
    "SciCoP": dict(folder="SciCoP", font="Times New Roman", size=10,
                   page=("Letter",), margin_cm=2.54),
}

A4 = (Cm(21.0), Cm(29.7))
LETTER = (Inches(8.5), Inches(11.0))


def setup(doc, prof):
    doc.styles["Normal"].font.name = prof["font"]
    doc.styles["Normal"].font.size = Pt(prof["size"])
    w, h = A4 if prof["page"][0] == "A4" else LETTER
    m = Cm(prof["margin_cm"])
    for sec in doc.sections:
        sec.page_width, sec.page_height = w, h
        sec.left_margin = sec.right_margin = m
        sec.top_margin = sec.bottom_margin = m


def _run(p, text, prof, size=None, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = prof["font"]
    r.font.size = Pt(size or prof["size"])
    r.bold = bold
    r.italic = italic
    return r


def title(doc, text, prof):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, prof, size=16, bold=True)
    p.paragraph_format.space_after = Pt(6)


def centered(doc, text, prof, size=None, italic=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, prof, size=size, italic=italic)
    p.paragraph_format.space_after = Pt(2)
    return p


def heading(doc, text, prof):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, prof, size=prof["size"] + 1, bold=True)


def para(doc, text, prof, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, prof)
    return p


def bold_label_para(doc, label, text, prof):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    _run(p, label + " ", prof, bold=True)
    _run(p, text, prof)
    return p


def perf_table(doc, r, prof):
    rows = [
        ("scalar (add i64)", r["overhead_ns"]["scalar_add_i64"]),
        ("float (mul f64)", r["overhead_ns"]["float_mul_f64"]),
        ("string (strlen)", r["overhead_ns"]["string_strlen"]),
        ("multi-return (divmod)", r["overhead_ns"]["multi_return_divmod"]),
        ("zero-copy buffer (fill)", r["overhead_ns"]["buffer_fill"]),
        ("ctypes baseline", r["ctypes_ns"]),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for c, txt in zip(hdr, ("Call kind", "Overhead (ns/call)")):
        c.paragraphs[0].add_run(txt).bold = True
    for label, ns in rows:
        cells = t.add_row().cells
        cells[0].text = label
        cells[1].text = f"{ns:.0f}"
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def figure(doc, path, caption, prof, width_in=5.3):
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = centered(doc, caption, prof, size=prof["size"] - 1, italic=True)
    cap.paragraph_format.space_after = Pt(8)


def build(name, prof, r):
    doc = Document()
    setup(doc, prof)

    title(doc, TITLE, prof)
    centered(doc, AUTHOR, prof)
    centered(doc, AFFIL, prof, size=prof["size"] - 1, italic=True)
    centered(doc, EMAIL, prof, size=prof["size"] - 1)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    speedup = r["speedup_vs_ctypes"]
    add_ns = r["overhead_ns"]["scalar_add_i64"]

    heading(doc, "Abstract", prof)
    para(doc,
         "Calling native code from Python is routine. The general-purpose "
         "bridges that make it convenient, ctypes and cffi, inspect argument "
         "types at call time, and that per-call cost dominates once the callee "
         "is small and called often. Nano-FFI moves the type dispatch to "
         "compile time instead. Given a Zig function and a declarative "
         "signature, its makeTrampoline runs Zig's comptime engine to generate "
         "a specialised wrapper whose argument unpacking compiles to "
         "straight-line code, with no runtime type inspection and no branch in "
         "the hot path. The bridge marshals "
         f"{r['n_types']} argument kinds, among them UTF-8 strings, raw bytes, "
         "zero-copy writable buffers, and multi-value tuple returns; Zig error "
         "unions map onto Python exceptions. On the reference platform a scalar "
         f"call costs about {add_ns:.0f} ns, roughly {speedup:.1f} times less "
         "than a minimal ctypes call. "
         f"{r['n_tests']} end-to-end tests and the Zig unit suite cover the "
         "surface. Nano-FFI is released under the MIT license with a frozen "
         "1.0 API.", prof)

    heading(doc, "1. Introduction", prof)
    para(doc,
         "Python's reach into native code rests on a few well-worn tools. "
         "ctypes and cffi are popular because they need no compiler at the call "
         "site: the caller describes a function's argument and return types, and "
         "the library marshals values accordingly at run time. That flexibility "
         "has a price. Every call walks a description of the signature, "
         "dispatches on each argument's type, and boxes or unboxes values "
         "through general-purpose paths. When the native function does real "
         "work, the overhead vanishes into the noise. When it is a handful of "
         "instructions in a tight loop, the marshalling is most of the cost.", prof)
    para(doc,
         "Compiled bridges avoid that by specialising ahead of time. Cython [1] "
         "compiles annotated Python to C; pybind11 and PyO3 generate C++ or Rust "
         "glue whose type handling is fixed at build time. They are powerful, "
         "general, and correspondingly large. Nano-FFI aims lower. Its one idea "
         "is that the wrapper for each function should be generated, fully typed, "
         "at compile time, in a language whose compile-time execution model makes "
         "that natural.", prof)
    para(doc,
         "That language is Zig [2]. Zig's comptime runs ordinary code during "
         "compilation, over ordinary values, including a function and a "
         "description of its signature. Nano-FFI uses this directly. The "
         "argument-unpacking loop is a comptime inline for over the signature, so "
         "the compiler emits explicit typed assignments rather than a runtime "
         "loop with per-element type tests. What Python calls into, then, is a "
         "dispatch path with no branch left to resolve.", prof)

    heading(doc, "2. Background", prof)
    para(doc,
         "A CPython extension module is a shared library that exposes an "
         "initialisation symbol and returns a module object populated with "
         "PyMethodDef entries. Inside those methods, values cross the boundary as "
         "PyObject pointers and convert to and from C types through the CPython "
         "C-API [3]. ctypes [4] and cffi [5] wrap that machinery behind a runtime "
         "type description, and the description is consulted on every call.", prof)
    para(doc,
         "Specialising removes that step. If the exact types are known when the "
         "wrapper is built, the conversions can be emitted as fixed code. Cython, "
         "pybind11 and PyO3 all do this, from different source languages. "
         "Nano-FFI's contribution is not to compete with them on breadth but to "
         "show how small and how fast the bridge becomes when the wrapper "
         "generator is itself a compile-time program written in Zig's comptime.", prof)

    heading(doc, "3. Design and implementation", prof)
    para(doc,
         "Nano-FFI holds one strict boundary. Exactly one source file, "
         "python_ext.zig, includes <Python.h>; every other module is pure Zig "
         "over C-ABI-compatible types. That isolation keeps the CPython "
         "dependency in one place and lets the core (the trampoline generator and "
         "the registry) be unit-tested without a Python interpreter running at "
         "all.", prof)
    bold_label_para(doc, "Comptime trampolines.",
         "The core is makeTrampoline(func, signature). It returns a wrapper of "
         "fixed C-ABI shape, fn (args: [*]const RawArg) callconv(.c) RawRet, that "
         "Python calls in place of func. The wrapper unpacks each argument at a "
         "comptime-known offset with an inline for over the signature. Because the "
         "loop is inline and the types are comptime, the compiler produces one "
         "typed assignment per argument, not a loop with runtime type checks. The "
         "return is handled the same way: pack picks the return conversion at "
         "compile time.", prof)
    bold_label_para(doc, "C-ABI marshalling.",
         "Values cross the internal boundary as an extern struct RawArg {tag, "
         "val}, where val is an extern union of the supported scalar "
         "representations plus a RawSlice {ptr, len} for variable-length data. "
         "Tagged unions are not C-ABI stable, so Nano-FFI carries an explicit tag "
         "alongside a C-compatible union instead. The return type RawRet holds the "
         "same union, an optional error pointer, and a small multi-return count.", prof)
    bold_label_para(doc, "Registry.",
         "A StringHashMap maps a name to a function pointer and its Signature. The "
         "dispatcher looks the entry up, checks arity, unpacks the Python tuple "
         "into a RawArg array, invokes the trampoline through the stored pointer, "
         "and converts the result back to a PyObject. The signature travels with "
         "the pointer, which is what later lets the module describe itself.", prof)

    heading(doc, "4. Type system and safety", prof)
    para(doc,
         "Nano-FFI marshals "
         f"{r['n_types']} argument kinds: " + ", ".join(r["supported_types"]) +
         ". Three choices keep the boundary safe, not just fast.", prof)
    para(doc,
         "Start with integer narrowing. A Python integer that will not fit the "
         "target width raises ValueError rather than truncating in silence. "
         "Unsigned targets reject negatives, and u64 uses the full unsigned "
         "range.", prof)
    para(doc,
         "Errors matter just as much. A wrapped Zig function may return an error "
         "union E!T. The trampoline unwraps it at compile time, so non-fallible "
         "functions keep the branch-free path, and on error it carries the Zig "
         "@errorName back. The Python layer raises that as a RuntimeError whose "
         "message is the error name.", prof)
    para(doc,
         "The third choice is about memory. A buffer argument is a writable "
         "Python buffer, a bytearray, a writable memoryview, or a NumPy array, "
         "handed to Zig as a mutable []u8 that points straight at Python's own "
         "storage. Nothing is copied. The buffer is acquired with "
         "PyObject_GetBuffer(PyBUF_WRITABLE) and released after the call; a "
         "read-only object is rejected with BufferError.", prof)
    para(doc,
         "The module is self-describing. list_functions() returns the registered "
         f"names ({r['n_functions']} built-in examples ship with the module), and "
         "signature(name) returns the argument names and types and the return "
         "type, with a list of types when the return is multi-value. PEP 561 "
         "stubs and a py.typed marker ship in the wheel.", prof)

    heading(doc, "5. Verification", prof)
    para(doc,
         "Correctness is checked at two levels. The pure-Zig core (trampoline "
         "generation, registry, allocator boundary) runs under Zig unit tests "
         "that need no Python interpreter, through a dedicated test root that "
         "leaves the <Python.h> module out. The full boundary is covered by "
         f"{r['n_tests']} end-to-end tests in CPython that call every built-in "
         "example across every type, and that includes the overflow, "
         "wrong-arity, empty-input, read-only-buffer, and divide-by-zero error "
         "paths. Both suites pass on the reference platform.", prof)

    heading(doc, "6. Performance", prof)
    para(doc,
         "I measure the wall-clock cost of a single call, averaged over "
         f"{r['iterations']:,} iterations after a warm-up, and report the median "
         f"of {r['repeats']} repeats to damp scheduler noise. The baseline is a "
         "minimal ctypes call into a trivial libc function (abs/fabs). That is a "
         "reference for dispatch overhead, not a like-for-like of the same "
         "computation. Every measurement uses a ReleaseFast build.", prof)
    para(doc,
         f"On CPython {r['python']} ({r['platform']}) the scalar add(i64, i64) "
         f"call costs about {r['overhead_ns']['scalar_add_i64']:.0f} ns. A float "
         f"multiply costs about {r['overhead_ns']['float_mul_f64']:.0f} ns, a "
         f"string-length call about {r['overhead_ns']['string_strlen']:.0f} ns, a "
         f"two-value divmod return about "
         f"{r['overhead_ns']['multi_return_divmod']:.0f} ns, and a zero-copy "
         f"buffer fill about {r['overhead_ns']['buffer_fill']:.0f} ns. The "
         f"minimal ctypes baseline is about {r['ctypes_ns']:.0f} ns, so a scalar "
         f"Nano-FFI call carries roughly {speedup:.1f} times less overhead. The "
         "variable-length and multi-value paths cost more than the scalar path, "
         "as one would expect, and both stay well under the general-purpose "
         "baseline.", prof)
    perf_table(doc, r, prof)
    figure(doc, os.path.join(FIG, "fig_overhead.png"),
           "Figure 1. Per-call overhead by argument kind.", prof)
    figure(doc, os.path.join(FIG, "fig_vs_ctypes.png"),
           "Figure 2. Scalar call overhead versus the ctypes baseline.", prof,
           width_in=3.6)

    heading(doc, "7. Discussion and limitations", prof)
    para(doc,
         "The benchmark measures dispatch overhead, not application throughput. "
         "For native functions that do substantial work the bridge cost is "
         "irrelevant, and Nano-FFI's advantage is confined to the small-callee, "
         "high-frequency regime it is built for. The comparison function is not "
         "the baseline's callee, so the ratio is best read as the cost of "
         "crossing the boundary, not a speedup of identical work. Two structural "
         "limits remain. The bridge caps arguments and return values at eight "
         "each, and a Zig-returned slice must be valid at the moment of return, "
         "whether static data or a view into the arguments; returning owned heap "
         "buffers is future work. The reported numbers are also single-platform. "
         "The build matrix covers Linux, Windows and macOS, but the quoted "
         "overhead is the Windows reference run.", prof)

    heading(doc, "8. Availability", prof)
    para(doc,
         "Nano-FFI is open source under the MIT license. The 1.0 release freezes "
         "the public API (call, version, list_functions, signature), the "
         "supported type names, and the exception mapping, and follows semantic "
         "versioning from there. Source, tests, the benchmark harness, and this "
         "manuscript's build scripts live in the repository. The extension builds "
         "with Zig 0.15.2 against CPython 3.10 or newer.", prof)

    heading(doc, "References", prof)
    for ref in [
        "[1] S. Behnel, R. Bradshaw, C. Citro, L. Dalcin, D. S. Seljebotn, K. "
        "Smith, \"Cython: The Best of Both Worlds,\" Computing in Science & "
        "Engineering, vol. 13, no. 2, pp. 31-39, 2011.",
        "[2] The Zig Programming Language. https://ziglang.org",
        "[3] Python/C API Reference Manual. https://docs.python.org/3/c-api/",
        "[4] ctypes - A foreign function library for Python. "
        "https://docs.python.org/3/library/ctypes.html",
        "[5] CFFI - C Foreign Function Interface for Python. "
        "https://cffi.readthedocs.io",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, ref, prof, size=prof["size"] - 1)

    out_dir = os.path.join(HERE, prof["folder"])
    os.makedirs(out_dir, exist_ok=True)
    out = os.path.join(out_dir, f"Carvajal_{name}_NanoFFI_manuscript.docx")
    doc.save(out)
    return out


def main(argv):
    r = compute_results()
    targets = argv[1:] if len(argv) > 1 else list(JOURNALS)
    for name in targets:
        if name not in JOURNALS:
            print(f"unknown journal: {name}; choices: {', '.join(JOURNALS)}")
            continue
        out = build(name, JOURNALS[name], r)
        print("wrote", out)


if __name__ == "__main__":
    main(sys.argv)
